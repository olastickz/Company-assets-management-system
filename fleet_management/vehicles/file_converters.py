"""
File converters for importing data from Excel and Word documents
"""

from openpyxl import load_workbook
from docx import Document
from datetime import datetime
import re
from io import StringIO
import csv


def parse_date(date_str):
    """Parse various date formats and return YYYY-MM-DD or None"""
    if not date_str or date_str.upper() == 'N/A':
        return None
    
    date_str = str(date_str).strip()
    
    # If it contains "2025" or earlier years, mark as pending
    if '2025' in date_str or '2024' in date_str:
        return None  # Will be marked as pending
    
    # Try various date formats
    formats = [
        '%d/%m/%Y',
        '%d-%m-%Y',
        '%d %b %Y',
        '%d %B %Y',
        '%Y-%m-%d',
        '%m/%d/%Y',
    ]
    
    for fmt in formats:
        try:
            parsed = datetime.strptime(date_str, fmt)
            return parsed.strftime('%Y-%m-%d')
        except ValueError:
            continue
    
    return None


def excel_to_csv_equipment(excel_file):
    """Convert Excel laptop tagging file to CSV format for equipment import
    
    CURRENT STRUCTURE (User-Reorganized):
    - iTECO: Headers row 3, data rows 4-28 (18 laptops)
      Columns: [empty, SUSIDIARY, ITEM DESCRIPTION, LOCATION, QTY, SERIAL NUMBER, YOP, TAG NUMBER, USER, REMARKS]
    
    - Sofworks: Headers row 4, data rows 5-19 (15 laptops)
      Columns: [empty, empty, ITEM DESCRIPTION, LOCATION, QTY, SERIAL NUMBER, YOP, TAG NUMBER, USER, REMARKS]
    
    - Telnet: Headers row 4, data rows 5-29 (25 laptops)
      Columns: [empty, empty, ITEM DESCRIPTION, LOCATION, QTY, SERIAL NUMBER, YOP, TAG NUMBER, USER, REMARKS]
    """
    
    wb = load_workbook(excel_file)
    equipment_data = []
    
    # Map sheet names to their exact configurations
    sheet_config = {
        'iteco': {
            'subsidiary': 'ITECO',
            'data_start_row': 4,
            'columns': {
                'subsidiary': 1,      # Col B has SUSIDIARY
                'description': 2,     # Col C
                'location': 3,        # Col D
                'qty': 4,             # Col E
                'serial_number': 5,   # Col F
                'yop': 6,             # Col G
                'tag_number': 7,      # Col H
                'user': 8,            # Col I
                'remarks': 9,         # Col J
            }
        },
        'sofworks': {
            'subsidiary': 'Softworks',
            'data_start_row': 5,
            'columns': {
                'description': 2,     # Col C (skip empty A & B)
                'location': 3,        # Col D
                'qty': 4,             # Col E
                'serial_number': 5,   # Col F
                'yop': 6,             # Col G
                'tag_number': 7,      # Col H
                'user': 8,            # Col I
                'remarks': 9,         # Col J
            }
        },
        'telnet': {
            'subsidiary': 'Telnet',
            'data_start_row': 5,
            'columns': {
                'description': 2,     # Col C (skip empty A & B)
                'location': 3,        # Col D
                'qty': 4,             # Col E
                'serial_number': 5,   # Col F
                'yop': 6,             # Col G
                'tag_number': 7,      # Col H
                'user': 8,            # Col I
                'remarks': 9,         # Col J
            }
        },
    }
    
    # Process each sheet
    for sheet_name in wb.sheetnames:
        sheet_name_lower = sheet_name.lower().strip()
        config = None
        
        # Match sheet to config
        for key, cfg in sheet_config.items():
            if key == sheet_name_lower:
                config = cfg
                break
        
        if not config:
            continue
        
        ws = wb[sheet_name]
        subsidiary = config['subsidiary']
        data_start_row = config['data_start_row']
        col_map = config['columns']
        
        # Process each data row
        for row_idx, row in enumerate(ws.iter_rows(min_row=data_start_row, values_only=True), start=data_start_row):
            if not any(row):
                continue
            
            # Extract values based on column positions
            description = row[col_map['description']] if col_map['description'] < len(row) and row[col_map['description']] else None
            location = row[col_map['location']] if col_map['location'] < len(row) and row[col_map['location']] else None
            qty = row[col_map['qty']] if col_map['qty'] < len(row) and row[col_map['qty']] else 1
            serial_number = row[col_map['serial_number']] if col_map['serial_number'] < len(row) and row[col_map['serial_number']] else None
            yop = row[col_map['yop']] if col_map['yop'] < len(row) and row[col_map['yop']] else None
            tag_number = row[col_map['tag_number']] if col_map['tag_number'] < len(row) and row[col_map['tag_number']] else None
            user = row[col_map['user']] if col_map['user'] < len(row) and row[col_map['user']] else None
            remarks = row[col_map['remarks']] if col_map['remarks'] < len(row) and row[col_map['remarks']] else None
            
            # Convert to strings and clean up
            description = str(description).strip() if description else None
            location = str(location).strip() if location else None
            serial_number = str(serial_number).strip() if serial_number else None
            tag_number = str(tag_number).strip() if tag_number else None
            user = str(user).strip() if user else None
            remarks = str(remarks).strip() if remarks else None
            
            # Skip empty rows or if no description or tag number
            if not description or not tag_number:
                continue
            
            # Skip header rows
            if description.lower() in ['item description', 'description']:
                continue
            
            # Format detailed description with all headers
            detailed_description = f"""Item Description: {description}
Location: {location or 'N/A'}
QTY: {qty or 'N/A'}
Serial Number: {serial_number or 'N/A'}
YOP (Year of Purchase): {yop if yop else 'null - to be added later'}
Tag Number: {tag_number or 'N/A'}
Assigned User: {user or 'Unassigned'}
Remarks: {remarks or 'N/A'}"""
            
            # Build equipment entry
            equipment_data.append({
                'name': description,
                'equipment_type': 'computer',
                'description': detailed_description,
                'location': location or 'Main Office',
                'status': 'active',
                'purchase_date': None,
                'cost': None,
                'notes': remarks if remarks else "",
                'subsidiary': subsidiary,
                'serial_number': serial_number,
                'tag_number': tag_number,
                'assigned_user': user,
                # Legacy field: kept for import compatibility only. Do NOT expose in any UI.
                'year_of_purchase': int(yop) if yop and str(yop).isdigit() else None,
                'quantity': int(qty) if qty and str(qty).isdigit() else 1,
                'remarks': remarks,
            })
    
    return equipment_data


def excel_to_csv_vehicles(excel_file):
    """Convert a generic vehicle Excel import file to CSV format."""

    wb = load_workbook(excel_file, data_only=True)
    vehicle_data = []
    # Use the first sheet by default
    ws = wb[wb.sheetnames[0]]

    header_row = None
    header_map = {}

    for row_idx, row in enumerate(ws.iter_rows(values_only=True), start=1):
        values = [str(cell).strip() if cell is not None else '' for cell in row]
        if not any(values):
            continue

        lower_values = [value.lower() for value in values]
        if any(key in value for value in lower_values for key in ('name', 'make', 'license_plate', 'vin', 'model', 'insurance', 'roadworthy', 'hackney')):
            header_row = row_idx
            header_map = {
                'name': next((i for i, value in enumerate(lower_values) if 'name' in value and 'vehicle' in value or value == 'name'), None),
                'make': next((i for i, value in enumerate(lower_values) if 'make' in value and 'model' not in value), None),
                'model': next((i for i, value in enumerate(lower_values) if 'model' in value), None),
                'vin_number': next((i for i, value in enumerate(lower_values) if 'vin' in value), None),
                'license_plate': next((i for i, value in enumerate(lower_values) if 'registration' in value or 'license plate' in value or 'reg no' in value or 'plate' == value), None),
                'vehicle_type': next((i for i, value in enumerate(lower_values) if 'vehicle type' in value or value == 'type'), None),
                'roadworthy_expiry': next((i for i, value in enumerate(lower_values) if 'roadworthy' in value), None),
                'hackney_permit': next((i for i, value in enumerate(lower_values) if 'hackney' in value), None),
                'license_expiry': next((i for i, value in enumerate(lower_values) if 'license expiry' in value or ('license' in value and 'expiry' in value)), None),
                'insurance_expiry': next((i for i, value in enumerate(lower_values) if 'insurance' in value), None),
            }
            break

    if header_row is None:
        raise ValueError('Excel file does not appear to contain a valid vehicle header row')

    for row in ws.iter_rows(min_row=header_row + 1, values_only=True):
        values = [str(cell).strip() if cell is not None else '' for cell in row]
        if not any(values):
            continue

        def get_value(key):
            idx = header_map.get(key)
            return values[idx] if idx is not None and idx < len(values) else ''

        name = get_value('name') or get_value('make')
        license_plate = get_value('license_plate')
        if not name or not license_plate:
            continue

        make = get_value('make')
        model = get_value('model')
        vin_number = get_value('vin_number')
        vehicle_type = get_value('vehicle_type') or 'car'
        roadworthy_date = parse_date(get_value('roadworthy_expiry'))
        hackney_date = parse_date(get_value('hackney_permit'))
        license_date = parse_date(get_value('license_expiry'))
        insurance_date = parse_date(get_value('insurance_expiry'))

        vehicle_data.append({
            'name': name,
            'make': make,
            'model': model,
            'vin_number': vin_number,
            'license_plate': license_plate,
            'vehicle_type': vehicle_type.lower(),
            'roadworthy_expiry': roadworthy_date,
            'hackney_permit': hackney_date,
            'license_expiry': license_date,
            'insurance_expiry': insurance_date,
        })

    return vehicle_data


def docx_to_csv_vehicles(docx_file):
    """Convert Word vehicle document to CSV format for vehicle import"""

    doc = Document(docx_file)
    vehicle_data = []

    if not doc.tables:
        return vehicle_data

    table = doc.tables[0]
    header_map = {}
    header_row_index = None

    # Attempt to detect the header row by matching common column names
    for idx, row in enumerate(table.rows[:3]):
        headers = [cell.text.strip().lower() for cell in row.cells]
        if any(h for h in headers if h):
            if any(keyword in h for h in headers for keyword in ('make', 'model', 'vin', 'registration', 'license', 'roadworthy', 'insurance')):
                header_row_index = idx
                header_map = {
                    'name': next((i for i, h in enumerate(headers) if 'name' in h and 'vehicle name' in h or h == 'name'), None),
                    'make': next((i for i, h in enumerate(headers) if 'make' in h and 'model' not in h), None),
                    'model': next((i for i, h in enumerate(headers) if 'model' in h), None),
                    'vin_number': next((i for i, h in enumerate(headers) if 'vin' in h), None),
                    'license_plate': next((i for i, h in enumerate(headers) if 'registration' in h or 'license plate' in h or 'reg no' in h), None),
                    'roadworthy_expiry': next((i for i, h in enumerate(headers) if 'roadworthy' in h), None),
                    'hackney_permit': next((i for i, h in enumerate(headers) if 'hackney' in h), None),
                    'license_expiry': next((i for i, h in enumerate(headers) if 'license expiry' in h or ('license' in h and 'expiry' in h)), None),
                    'insurance_expiry': next((i for i, h in enumerate(headers) if 'insurance' in h), None),
                    'vehicle_type': next((i for i, h in enumerate(headers) if 'vehicle type' in h or h == 'type'), None),
                }
                break

    if header_row_index is None:
        header_row_index = 0
        header_map = {
            'name': 1,
            'make': 1,
            'model': None,
            'vin_number': None,
            'license_plate': 2,
            'roadworthy_expiry': 3,
            'hackney_permit': 4,
            'license_expiry': 5,
            'insurance_expiry': 6,
            'vehicle_type': None,
        }

    for row in table.rows[header_row_index + 1:]:
        cells = [cell.text.strip() for cell in row.cells]
        if not any(cells):
            continue

        def get_cell_value(key):
            idx = header_map.get(key)
            return cells[idx].strip() if idx is not None and idx < len(cells) else ''

        name = get_cell_value('name') or get_cell_value('make')
        make = get_cell_value('make')
        model = get_cell_value('model')
        vin_number = get_cell_value('vin_number')
        license_plate = get_cell_value('license_plate')
        vehicle_type = get_cell_value('vehicle_type')
        if not vehicle_type:
            vehicle_type = 'van' if make and ('van' in make.lower() or 'hiace' in make.lower()) else 'car'

        if not name or not license_plate:
            continue

        roadworthy_date = parse_date(get_cell_value('roadworthy_expiry'))
        hackney_date = parse_date(get_cell_value('hackney_permit'))
        license_date = parse_date(get_cell_value('license_expiry'))
        insurance_date = parse_date(get_cell_value('insurance_expiry'))

        vehicle_data.append({
            'name': name,
            'make': make,
            'model': model,
            'vin_number': vin_number,
            'license_plate': license_plate,
            'vehicle_type': vehicle_type.lower() if vehicle_type else 'car',
            'roadworthy_expiry': roadworthy_date,
            'hackney_permit': hackney_date,
            'license_expiry': license_date,
            'insurance_expiry': insurance_date,
        })

    return vehicle_data


def vehicles_to_csv_string(vehicle_list):
    """Convert vehicle data list to CSV string"""
    output = StringIO()
    writer = csv.DictWriter(
        output,
        fieldnames=['name', 'license_plate', 'vin_number', 'make', 'model', 'vehicle_type', 'roadworthy_expiry', 'hackney_permit', 'license_expiry', 'insurance_expiry']
    )
    writer.writeheader()
    writer.writerows(vehicle_list)
    return output.getvalue()


def equipment_to_csv_string(equipment_list):
    """Convert equipment data list to CSV string"""
    output = StringIO()
    writer = csv.DictWriter(
        output,
        # Include `year_of_purchase` for backward import compatibility only.
        # UI/forms should not implement or display this field.
        fieldnames=['name', 'equipment_type', 'description', 'location', 'status', 'purchase_date', 'cost', 'notes', 'subsidiary', 'serial_number', 'tag_number', 'assigned_user', 'year_of_purchase', 'quantity', 'remarks']
    )
    writer.writeheader()
    writer.writerows(equipment_list)
    return output.getvalue()
